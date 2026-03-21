"""
core/docx_utils.py — وثائق Word احترافية بـ RTL عربي مضمون 100%
كل فقرة وكل run يحمل w:bidi + w:rtl + w:cs مباشرةً

v6: إضافة logging + error handling احترافي لكل نموذج
"""
import io
import logging
from pathlib import Path
from datetime import date

from django.conf import settings
from django.http import HttpResponse

logger = logging.getLogger(__name__)

MAROON   = "8A1538"
MAROON_R = (138, 21, 56)
LIGHT_BG = "FDF2F5"
FONT_AR  = "Arial"   # يدعم العربية في كل Word إصدارات

_ASSETS = Path(__file__).parent / "assets"


# ════════════════════════════════════════════════════════════════════
# primitive helpers
# ════════════════════════════════════════════════════════════════════

def _read_img(name):
    for p in [_ASSETS / name,
               Path(settings.BASE_DIR) / "static" / "images" / name,
               Path(settings.BASE_DIR) / "static" / name]:
        if p.exists():
            return p.read_bytes()
    return None

def _pt(n):
    from docx.shared import Pt;  return Pt(n)
def _cm(n):
    from docx.shared import Cm;  return Cm(n)
def _rgb(r, g, b):
    from docx.shared import RGBColor;  return RGBColor(r, g, b)
def _el(tag):
    from docx.oxml import OxmlElement;  return OxmlElement(tag)
def _qn(tag):
    from docx.oxml.ns import qn;  return qn(tag)


# ────────────────────────────────────────────────────────────────────
# RTL على مستوى الفقرة — الطريقة الوحيدة الموثوقة في python-docx
# ────────────────────────────────────────────────────────────────────

def _ppr_rtl(para, align="right", sb=60, sa=60):
    """يضبط <w:pPr> للفقرة: bidi + jc + spacing"""
    pPr = para._p.get_or_add_pPr()

    # w:bidi — يجعل الفقرة RTL
    bidi = _el("w:bidi")
    pPr.insert(0, bidi)

    # w:jc — المحاذاة
    jc = _el("w:jc")
    jc.set(_qn("w:val"), align)
    pPr.append(jc)

    # w:spacing
    sp = _el("w:spacing")
    sp.set(_qn("w:before"), str(sb))
    sp.set(_qn("w:after"),  str(sa))
    pPr.append(sp)

    return pPr


def _rpr_arabic(run, size_pt=11, bold=False, color_rgb=None):
    """
    يضبط <w:rPr> للـ run:
      - w:rFonts مع w:cs='Arial'  ← يجبر Word على الخط العربي
      - w:rtl                      ← يجعل الـ run RTL
      - w:sz / w:szCs              ← حجم الخط للنصوص العربية
    """
    run.bold      = bold
    run.font.size = _pt(size_pt)
    if color_rgb:
        run.font.color.rgb = _rgb(*color_rgb)

    rPr = run._r.get_or_add_rPr()

    # rFonts
    rF = _el("w:rFonts")
    rF.set(_qn("w:ascii"),    FONT_AR)
    rF.set(_qn("w:hAnsi"),   FONT_AR)
    rF.set(_qn("w:cs"),      FONT_AR)   # ← هذا هو الأهم للعربية
    rF.set(_qn("w:eastAsia"), FONT_AR)
    rPr.insert(0, rF)

    # rtl run
    rPr.append(_el("w:rtl"))

    # حجم الخط بـ half-points (للـ CS script)
    sz   = _el("w:sz");   sz.set(_qn("w:val"),   str(int(size_pt * 2)));  rPr.append(sz)
    szCs = _el("w:szCs"); szCs.set(_qn("w:val"), str(int(size_pt * 2)));  rPr.append(szCs)

    return rPr


def _add_rtl_para(doc, text, size=11, bold=False, color=None,
                  align="right", sb=60, sa=60):
    """أضف فقرة عربية RTL كاملة"""
    p = doc.add_paragraph()
    _ppr_rtl(p, align=align, sb=sb, sa=sa)
    run = p.add_run(text or "")
    _rpr_arabic(run, size_pt=size, bold=bold, color_rgb=color)
    return p


# ────────────────────────────────────────────────────────────────────
# الجداول: bidiVisual + كل خلية RTL
# ────────────────────────────────────────────────────────────────────

def _tbl_rtl(tbl):
    """يضبط الجدول على RTL بصري"""
    tblPr = tbl._tbl.find(_qn("w:tblPr"))
    if tblPr is None:
        tblPr = _el("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    bv = _el("w:bidiVisual")
    tblPr.append(bv)
    jc = _el("w:jc")
    jc.set(_qn("w:val"), "right")
    tblPr.append(jc)


def _cell_rtl(cell, text="", size=10, bold=False, color=None,
              bg=None, border_color=None, align="right"):
    """يضبط الخلية + فقرتها على RTL"""
    # خلفية
    if bg:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = _el("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), bg)
        tcPr.append(shd)

    # حدود
    if border_color:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBdr = _el("w:tcBdr")
        for side in ("top", "left", "bottom", "right"):
            el = _el(f"w:{side}")
            el.set(_qn("w:val"), "single")
            el.set(_qn("w:sz"), "4")
            el.set(_qn("w:color"), border_color)
            tcBdr.append(el)
        tcPr.append(tcBdr)

    # الفقرة
    p = cell.paragraphs[0]
    _ppr_rtl(p, align=align, sb=40, sa=40)
    p.paragraph_format.right_indent = _cm(0.2)
    p.paragraph_format.left_indent  = _cm(0.1)

    if text:
        run = p.add_run(str(text))
        _rpr_arabic(run, size_pt=size, bold=bold, color_rgb=color)

    return p


def _clear_borders(cell):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcBdr = _el("w:tcBdr")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _el(f"w:{s}")
        el.set(_qn("w:val"), "none")
        tcBdr.append(el)
    tcPr.append(tcBdr)


# ════════════════════════════════════════════════════════════════════
# عناصر التصميم
# ════════════════════════════════════════════════════════════════════

def _new_doc():
    from docx import Document
    doc = Document()

    # A4 + هوامش
    for sec in doc.sections:
        sec.page_width    = _cm(21)
        sec.page_height   = _cm(29.7)
        sec.right_margin  = _cm(2.2)
        sec.left_margin   = _cm(2.2)
        sec.top_margin    = _cm(3.8)
        sec.bottom_margin = _cm(3.2)

    # هيدر + فوتر
    _attach_header_footer(doc)
    return doc


def _attach_header_footer(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sec = doc.sections[0]

    def _img_para(container, img_bytes):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.clear()
        pPr = p._p.get_or_add_pPr()
        sp = _el("w:spacing")
        sp.set(_qn("w:before"), "0")
        sp.set(_qn("w:after"), "0")
        pPr.append(sp)
        p.add_run().add_picture(io.BytesIO(img_bytes), width=_cm(16.6))

    hdr = _read_img("pdf_header.jpg") or _read_img("pdf_header.png")
    if hdr:
        sec.header_distance = _cm(0)
        sec.header.is_linked_to_previous = False
        _img_para(sec.header, hdr)

    ftr = _read_img("pdf_footer.jpg") or _read_img("pdf_footer.png")
    if ftr:
        sec.footer_distance = _cm(0)
        sec.footer.is_linked_to_previous = False
        _img_para(sec.footer, ftr)


def _title_block(doc, title, subtitle=""):
    """عنوان النموذج المركزي + خط مزدوج"""
    _add_rtl_para(doc, " ", size=6, sb=0, sa=0)
    _add_rtl_para(doc, title, size=18, bold=True,
                  color=MAROON_R, align="center", sb=0, sa=4)
    if subtitle:
        _add_rtl_para(doc, subtitle, size=10,
                      color=(90, 90, 90), align="center", sb=0, sa=6)
    # خط مزدوج
    _divider(doc, style="double", sz="12")


def _divider(doc, style="single", sz="6", color=MAROON):
    p = doc.add_paragraph()
    _ppr_rtl(p, sb=2, sa=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = _el("w:pBdr")
    bot  = _el("w:bottom")
    bot.set(_qn("w:val"),   style)
    bot.set(_qn("w:sz"),    sz)
    bot.set(_qn("w:color"), color)
    bot.set(_qn("w:space"), "1")
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_title(doc, title):
    """عنوان قسم: شريط يميني + خلفية"""
    p = doc.add_paragraph()
    _ppr_rtl(p, sb=80, sa=20)
    pPr = p._p.get_or_add_pPr()

    # شريط يميني ملوّن
    pBdr = _el("w:pBdr")
    right = _el("w:right")
    right.set(_qn("w:val"),   "single")
    right.set(_qn("w:sz"),    "28")
    right.set(_qn("w:color"), MAROON)
    right.set(_qn("w:space"), "6")
    pBdr.append(right)
    pPr.append(pBdr)

    # خلفية
    shd = _el("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), LIGHT_BG)
    pPr.append(shd)

    run = p.add_run("  " + title)
    _rpr_arabic(run, size_pt=12, bold=True, color_rgb=MAROON_R)


def _info_grid(doc, pairs):
    """
    جدول بيانات احترافي: تسمية | قيمة | تسمية | قيمة
    pairs = [(lbl1, val1, lbl2, val2), ...]
    """
    tbl = doc.add_table(rows=len(pairs), cols=4)
    _tbl_rtl(tbl)
    tbl.style = "Table Grid"
    col_w = [_cm(4.7), _cm(2.8), _cm(4.7), _cm(2.8)]

    for i, row_data in enumerate(pairs):
        for j, cell in enumerate(tbl.rows[i].cells):
            cell.width = col_w[j]
            _clear_borders(cell)
            is_lbl = (j % 2 == 1)
            bg = "F0E0E8" if is_lbl else "FAFAFA"
            _cell_rtl(cell, text=str(row_data[j] or ""), size=10,
                      bold=is_lbl, color=MAROON_R if is_lbl else None, bg=bg)

    _sp(doc, 6)


def _content_box(doc, content, bg="FFFBFC"):
    """صندوق نص بحد يميني ملوّن"""
    tbl  = doc.add_table(rows=1, cols=1)
    _tbl_rtl(tbl)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    cell.width = _cm(16.6)

    # حدود الصندوق
    tcPr  = cell._tc.get_or_add_tcPr()
    tcBdr = _el("w:tcBdr")
    borders = {"right": (MAROON, "10"), "top": ("E5E7EB","4"),
               "bottom": ("E5E7EB","4"), "left": ("E5E7EB","4")}
    for side, (col, sz) in borders.items():
        el = _el(f"w:{side}")
        el.set(_qn("w:val"), "single")
        el.set(_qn("w:sz"), sz)
        el.set(_qn("w:color"), col)
        tcBdr.append(el)
    tcPr.append(tcBdr)

    # خلفية
    shd = _el("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), bg)
    tcPr.append(shd)

    p = cell.paragraphs[0]
    _ppr_rtl(p, sb=60, sa=60)
    p.paragraph_format.right_indent = _cm(0.4)
    p.paragraph_format.left_indent  = _cm(0.3)
    pPr = p._p.get_or_add_pPr()
    lsp = _el("w:spacing")
    lsp.set(_qn("w:line"), "360")   # line spacing 1.5
    lsp.set(_qn("w:lineRule"), "auto")
    pPr.append(lsp)

    run = p.add_run(content or "")
    _rpr_arabic(run, size_pt=11)
    _sp(doc, 6)


def _sig_table(doc, labels):
    """جدول توقيعات أفقي RTL"""
    n   = len(labels)
    tbl = doc.add_table(rows=3, cols=n)
    _tbl_rtl(tbl)
    tbl.style = "Table Grid"
    w = _cm(16.6 / n)

    for j, lbl in enumerate(labels):
        # ─ صف مسافة التوقيع ─
        c0 = tbl.rows[0].cells[j]
        c0.width = w
        _clear_borders(c0)
        shd = _el("w:shd"); shd.set(_qn("w:val"),"clear")
        shd.set(_qn("w:color"),"auto"); shd.set(_qn("w:fill"),"FAFAFA")
        c0._tc.get_or_add_tcPr().append(shd)
        p0 = c0.paragraphs[0]
        _ppr_rtl(p0, align="center", sb=200, sa=0)

        # ─ صف خط التوقيع ─
        c1 = tbl.rows[1].cells[j]
        c1.width = w
        _clear_borders(c1)
        tcPr  = c1._tc.get_or_add_tcPr()
        shd1  = _el("w:shd"); shd1.set(_qn("w:val"),"clear")
        shd1.set(_qn("w:color"),"auto"); shd1.set(_qn("w:fill"),"FAFAFA")
        tcPr.append(shd1)
        tcBdr = _el("w:tcBdr")
        top   = _el("w:top")
        top.set(_qn("w:val"),"single"); top.set(_qn("w:sz"),"10")
        top.set(_qn("w:color"),"444444"); tcBdr.append(top)
        tcPr.append(tcBdr)
        p1 = c1.paragraphs[0]
        _ppr_rtl(p1, align="center", sb=0, sa=20)

        # ─ صف اسم الموقّع ─
        c2 = tbl.rows[2].cells[j]
        c2.width = w
        _clear_borders(c2)
        shd2 = _el("w:shd"); shd2.set(_qn("w:val"),"clear")
        shd2.set(_qn("w:color"),"auto"); shd2.set(_qn("w:fill"),"F5E8EC")
        c2._tc.get_or_add_tcPr().append(shd2)
        p2 = c2.paragraphs[0]
        _ppr_rtl(p2, align="center", sb=30, sa=30)
        run = p2.add_run(lbl)
        _rpr_arabic(run, size_pt=10, bold=True, color_rgb=MAROON_R)

    _sp(doc, 8)


def _notice(doc, text):
    _divider(doc, sz="4", color="CCCCCC")
    _add_rtl_para(doc, text, size=8, color=(170,170,170), align="center", sb=20, sa=0)


def _sp(doc, pts=6):
    p = doc.add_paragraph()
    _ppr_rtl(p, sb=0, sa=pts)


def _docx_resp(doc, fn):
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    resp["Content-Disposition"] = f'attachment; filename="{fn}"'
    return resp


# ════════════════════════════════════════════════════════════════════
# نموذج 1 — إنذار سلوكي
# ════════════════════════════════════════════════════════════════════

def generate_warning_docx(ctx):
    try:
        inf   = ctx["infraction"]
        sname = getattr(ctx.get("school"), "name",
                        "مدرسة الشحانية الإعدادية الثانوية للبنين")
        doc = _new_doc()

        _title_block(doc, "نموذج إنذار سلوكي", sname)

        _section_title(doc, "بيانات الطالب")
        _info_grid(doc, [
            ("اسم الطالب:",  inf.student.full_name,
             "الرقم الشخصي:", str(inf.student.username or "")),
            ("الصف / الفصل:", str(ctx.get("class_name") or "___________"),
             "التاريخ:", str(inf.date)),
        ])

        _section_title(doc, "نوع المخالفة السلوكية")
        cat_text = (
            f"[{inf.violation_category.code}]  {inf.violation_category.name_ar}"
            f"  |  الدرجة: {inf.get_level_display()}"
            if inf.violation_category
            else (inf.description or "")[:120]
        )
        _content_box(doc, cat_text, bg="FFF0F3")

        _section_title(doc, "تفاصيل الواقعة")
        _content_box(doc, inf.description or "___________")

        _section_title(doc, "الإجراء المتخذ وفق اللوائح")
        action = (getattr(inf, "action_taken", None) or
                  getattr(getattr(inf, "violation_category", None),
                          "default_action", None) or "___________")
        _content_box(doc, action)

        _sp(doc, 8)
        count = ctx.get("infraction_count", 1)
        order = "أول" if count == 1 else "ثانٍ" if count == 2 else "نهائي"
        p = doc.add_paragraph()
        _ppr_rtl(p, sb=0, sa=60)
        p.paragraph_format.right_indent = _cm(0.3)
        run = p.add_run(
            f"هذا إنذار سلوكي {order} للطالب المذكور أعلاه. "
            "يُرجى الالتزام التام بلوائح السلوك والانضباط المدرسية.")
        _rpr_arabic(run, size_pt=11, bold=True)

        _sp(doc, 10)
        _section_title(doc, "التوقيعات")
        _sig_table(doc, ["رائد الفصل", "منسق السلوك", "وكيل شؤون الطلبة"])
        _sig_table(doc, ["توقيع الطالب", "توقيع ولي الأمر", "تاريخ التسليم: ____/__/__"])
        _notice(doc, f"SchoolOS v6  |  {ctx.get('generated_at', date.today())}")

        return _docx_resp(doc, f"warning_{inf.student.username}_{inf.date}.docx")

    except ImportError as e:
        logger.error("python-docx غير مثبت: %s", e)
        return HttpResponse("مكتبة Word غير متاحة — pip install python-docx", status=500)
    except KeyError as e:
        logger.error("بيانات ناقصة في generate_warning_docx: %s", e)
        return HttpResponse(f"بيانات ناقصة: {e}", status=400)
    except Exception as e:
        logger.exception("خطأ غير متوقع في generate_warning_docx: %s", e)
        return HttpResponse("حدث خطأ أثناء إنشاء الوثيقة", status=500)


# ════════════════════════════════════════════════════════════════════
# نموذج 2 — تعهد ولي الأمر
# ════════════════════════════════════════════════════════════════════

def generate_parent_undertaking_docx(ctx):
    try:
        inf   = ctx["infraction"]
        sname = getattr(ctx.get("school"), "name",
                        "مدرسة الشحانية الإعدادية الثانوية للبنين")
        doc = _new_doc()

        _title_block(doc, "تعهد ولي الأمر", sname)

        _section_title(doc, "بيانات ولي الأمر")
        _info_grid(doc, [
            ("اسم ولي الأمر:", ctx.get("parent_name") or "___________",
             "رقم الهوية:",    ctx.get("parent_id")   or "___________"),
            ("رقم الهاتف:",   ctx.get("parent_phone") or "___________",
             "البريد الإلكتروني:", ctx.get("parent_email") or "___________"),
        ])

        _section_title(doc, "بيانات الطالب")
        _info_grid(doc, [
            ("اسم الطالب:",  inf.student.full_name,
             "الرقم الشخصي:", str(inf.student.username or "")),
            ("الصف / الفصل:", str(ctx.get("class_name") or "___________"),
             "تاريخ المخالفة:", str(inf.date)),
        ])

        _section_title(doc, "المخالفة المُبلَّغ عنها")
        vtext = (f"[{inf.violation_category.code}]  {inf.violation_category.name_ar}"
                 if inf.violation_category else (inf.description or "")[:100])
        _content_box(doc, vtext, bg="FFF0F3")

        _section_title(doc, "نص التعهد")
        _content_box(doc,
            "أنا ولي أمر الطالب المذكور أعلاه، أُقرّ باطلاعي على المخالفة السلوكية الصادرة بحق ابني، "
            "وأتعهد بمتابعته وتوجيهه نحو الالتزام بلوائح المدرسة وأنظمتها، وعدم تكرار هذا السلوك "
            "مستقبلاً، والتعاون الكامل مع إدارة المدرسة في كل ما يخدم مصلحة ابني التعليمية والسلوكية.")

        _sp(doc, 10)
        _sig_table(doc, ["توقيع ولي الأمر", "وكيل شؤون الطلبة", "ختم المدرسة"])
        _notice(doc, f"SchoolOS v6  |  {ctx.get('generated_at', date.today())}")

        return _docx_resp(doc, f"parent_undertaking_{inf.student.username}.docx")

    except ImportError as e:
        logger.error("python-docx غير مثبت: %s", e)
        return HttpResponse("مكتبة Word غير متاحة — pip install python-docx", status=500)
    except KeyError as e:
        logger.error("بيانات ناقصة في generate_parent_undertaking_docx: %s", e)
        return HttpResponse(f"بيانات ناقصة: {e}", status=400)
    except Exception as e:
        logger.exception("خطأ غير متوقع في generate_parent_undertaking_docx: %s", e)
        return HttpResponse("حدث خطأ أثناء إنشاء الوثيقة", status=500)


# ════════════════════════════════════════════════════════════════════
# نموذج 3 — تعهد الطالب
# ════════════════════════════════════════════════════════════════════

def generate_student_undertaking_docx(ctx):
    try:
        inf   = ctx["infraction"]
        sname = getattr(ctx.get("school"), "name",
                        "مدرسة الشحانية الإعدادية الثانوية للبنين")
        doc = _new_doc()

        _title_block(doc, "تعهد الطالب", sname)

        _section_title(doc, "بيانات الطالب")
        _info_grid(doc, [
            ("اسم الطالب:",  inf.student.full_name,
             "الرقم الشخصي:", str(inf.student.username or "")),
            ("الصف / الفصل:", str(ctx.get("class_name") or "___________"),
             "التاريخ:", str(inf.date)),
        ])

        _section_title(doc, "المخالفة المرتكبة")
        vtext = (f"[{inf.violation_category.code}]  {inf.violation_category.name_ar}"
                 if inf.violation_category else (inf.description or "")[:100])
        _content_box(doc, vtext, bg="FFF0F3")

        _section_title(doc, "نص التعهد")
        _content_box(doc,
            "أنا الطالب المذكور أعلاه، أُقرّ باطلاعي على المخالفة السلوكية الصادرة بحقي، "
            "وأتعهد بعدم تكرار هذا السلوك مستقبلاً، والالتزام التام بلوائح المدرسة وتعليماتها، "
            "وأن أكون قدوة حسنة لزملائي في الالتزام والانضباط، مدركاً أن تكرار هذا السلوك "
            "سيعرضني لعقوبات أشد وفق لائحة السلوك المعتمدة.")

        _sp(doc, 10)
        _sig_table(doc, ["توقيع الطالب", "توقيع ولي الأمر", "رائد الفصل"])
        _notice(doc, f"SchoolOS v6  |  {ctx.get('generated_at', date.today())}")

        return _docx_resp(doc, f"student_undertaking_{inf.student.username}.docx")

    except ImportError as e:
        logger.error("python-docx غير مثبت: %s", e)
        return HttpResponse("مكتبة Word غير متاحة — pip install python-docx", status=500)
    except KeyError as e:
        logger.error("بيانات ناقصة في generate_student_undertaking_docx: %s", e)
        return HttpResponse(f"بيانات ناقصة: {e}", status=400)
    except Exception as e:
        logger.exception("خطأ غير متوقع في generate_student_undertaking_docx: %s", e)
        return HttpResponse("حدث خطأ أثناء إنشاء الوثيقة", status=500)


# ════════════════════════════════════════════════════════════════════
# وثيقة 4 — لائحة السلوك الكاملة
# ════════════════════════════════════════════════════════════════════

def generate_policy_docx(ctx):
    try:
     return _generate_policy_docx_inner(ctx)
    except ImportError as e:
        logger.error("python-docx غير مثبت: %s", e)
        return HttpResponse("مكتبة Word غير متاحة — pip install python-docx", status=500)
    except Exception as e:
        logger.exception("خطأ غير متوقع في generate_policy_docx: %s", e)
        return HttpResponse("حدث خطأ أثناء إنشاء الوثيقة", status=500)


def _generate_policy_docx_inner(ctx):
    doc = _new_doc()
    yr  = ctx.get("academic_year", "2025-2026")

    _title_block(doc, "لائحة السلوك والانضباط المدرسية",
                 f"مدرسة الشحانية الإعدادية الثانوية للبنين  |  {yr}  |  الإصدار 1.0")

    # ── أقسام نصية ──
    sections_text = [
        ("1. الهدف من اللائحة",
         "تهدف هذه اللائحة إلى توفير بيئة تعليمية آمنة ومنتجة تعزز قيم الانضباط والاحترام المتبادل، "
         "استناداً إلى: الاستراتيجية الوطنية للتعليم 2024–2030، قانون 9/2017 (التعليم المدرسي)، "
         "قانون 25/2001 (الأحداث الجانحون)، ولوائح وزارة التربية والتعليم 2025."),
        ("2. مجال التطبيق",
         "تسري هذه اللائحة على جميع طلبة المدرسة في كافة الفضاءات المدرسية والأنشطة المرتبطة بها."),
        ("3. مبادئ التعامل السلوكي",
         "الإيجابيات: بناء علاقات إيجابية وتعزيز السلوك الجيد.\n"
         "التصحيحيات: معالجة السلوك السلبي بأسلوب تصاعدي يراعي الظروف الفردية."),
    ]
    for title, content in sections_text:
        _section_title(doc, title)
        _content_box(doc, content)

    # ── 4. جدول المخالفات ──
    _section_title(doc, "4. جدول المخالفات السلوكية (A / B / C / D)")
    _sp(doc, 6)

    categories = [
        ("A", MAROON,   "مخالفات بسيطة",        "تنبيه / إنذار كتابي",          "D1FAE5",
         [("A1","غياب بدون عذر"),("A2","تأخر متكرر"),("A3","إخلال بالنظام"),
          ("A4","عدم إحضار الأدوات"),("A5","إهمال في المظهر")]),
        ("B","92400E",  "مخالفات متوسطة",        "إنذار + استدعاء ولي الأمر",    "FEF3C7",
         [("B1","الغش في الاختبارات"),("B2","التنمر"),("B3","إتلاف الممتلكات"),
          ("B4","الخروج بدون إذن"),("B5","التدخين")]),
        ("C","991B1B",  "مخالفات جسيمة",         "فصل مؤقت + لجنة انضباط",       "FEE2E2",
         [("C1","الاعتداء الجسدي"),("C2","التهديد"),("C3","السرقة"),
          ("C4","التزوير"),("C5","الاعتداء الإلكتروني")]),
        ("D","1F2937",  "مخالفات بالغة الخطورة", "فصل نهائي + إحالة للجهات",     "F3F4F6",
         [("D1","إحضار أسلحة"),("D2","الاتجار بالمخدرات"),("D3","الاعتداء الجنسي"),
          ("D4","إشعال الحرائق"),("D5","تهديد الأمن المدرسي")]),
    ]

    for cat_code, hdr_col, label, penalty, row_bg, items in categories:
        tbl = doc.add_table(rows=len(items) + 1, cols=3)
        _tbl_rtl(tbl)
        tbl.style = "Table Grid"
        cw = [_cm(1.8), _cm(9.8), _cm(5.0)]

        # رأس
        for j, (cell, txt) in enumerate(zip(tbl.rows[0].cells,
                                             [f"الفئة {cat_code}", label, penalty])):
            cell.width = cw[j]
            _cell_rtl(cell, txt, size=10, bold=True,
                      color=(255,255,255), bg=hdr_col, border_color=hdr_col)

        # صفوف
        for i, (code, name) in enumerate(items, 1):
            bg = row_bg if i % 2 == 1 else "FFFFFF"
            for j, (cell, txt) in enumerate(zip(
                    tbl.rows[i].cells, [code, name, penalty.split("+")[0].strip()])):
                cell.width = cw[j]
                _cell_rtl(cell, txt, size=10, bold=(j == 0),
                          bg=bg, border_color="E5E7EB",
                          align="center" if j == 0 else "right")

        _sp(doc, 8)

    # ── 5. سلّم العقوبات ──
    _section_title(doc, "5. سلّم العقوبات التدريجي")
    scale = [
        ("المستوى 1", "تنبيه شفهي / كتابي / اتصال هاتفي",      "D1FAE5"),
        ("المستوى 2", "إنذار أول + استدعاء ولي الأمر + تعهد",   "FEF3C7"),
        ("المستوى 3", "إنذار ثانٍ + خطة سلوكية + أخصائي",       "FEE2E2"),
        ("المستوى 4", "إنذار نهائي + لجنة انضباط + فصل مؤقت",   "FECACA"),
        ("المستوى 5", "فصل نهائي / إحالة للجهات المختصة",        "FCA5A5"),
    ]
    stbl = doc.add_table(rows=5, cols=2)
    _tbl_rtl(stbl)
    stbl.style = "Table Grid"
    for i, (lvl, desc, bg) in enumerate(scale):
        stbl.rows[i].cells[0].width = _cm(3.0)
        stbl.rows[i].cells[1].width = _cm(13.6)
        _cell_rtl(stbl.rows[i].cells[0], lvl, size=10, bold=True,
                  bg=bg, border_color="E5E7EB")
        _cell_rtl(stbl.rows[i].cells[1], desc, size=10,
                  bg=bg, border_color="E5E7EB")

    # ── التوقيعات ──
    _sp(doc, 14)
    _divider(doc, "double", "12")
    _sig_table(doc, ["مدير المدرسة", "وكيل شؤون الطلبة", "مسؤول الجودة"])
    _notice(doc, f"SchoolOS v6  |  {ctx.get('generated_at', date.today())}  |  الإصدار 1.0")

    return _docx_resp(doc, "behavior_policy_2025-2026.docx")